from selenium import webdriver
from bs4 import BeautifulSoup
import time 
import random
import docx
# import re
import datetime
# abc=re.compile("[a-zA-Z]")
UU= "http://www.csrc.gov.cn/pub/zjhpublic/"

bb= webdriver.Firefox()
bb.get('http://www.csrc.gov.cn/pub/zjhpublic/index.htm?channel=3300/3313')

bb.switch_to_default_content()
frame = bb.find_elements_by_tag_name('iframe')[0]
bb.switch_to_frame(frame)

#x写入url地址
def url_to(list):
    with open("zjh2_url.txt",'a',encoding='utf-8') as ff:              #换
        ff.writelines(datetime.datetime.now().strftime('%Y/%m/%d %H:%m')+"----------------------------------------------------"+"\n")	
        for i in list:
	        ff.writelines(i+"\n")

def geturl(bb,alist):
    htm=BeautifulSoup(bb.page_source,'lxml')
    aa=htm.select('#documentContainer .row .mc a')
    for ii in aa:
        hr1 = ii.get("href").replace("../","")
        href = UU+hr1
        alist.append(href)


def tlist_to(alist,zlbb):
    for index,hr in enumerate(alist):
        try:
            zlbb.get(hr)
        except:
            print("一个详细页出错")
            time.sleep(2)
            continue
        time.sleep(random.randint(1,3))      #sleep
        html= BeautifulSoup(zlbb.page_source,features='lxml')
        try:
            ti = html.select(".headInfo #lTitle")[0].text.strip()                     #要换
        except:
            ti = str(index)+"title"
        try:
            tim = zlbb.find_element_by_xpath("/html/body/div/div[2]/div[1]/table/tbody/tr[2]/td/table/tbody/tr/td[2]/span").text.strip()                    #要换
        except:
            tim = str(index)+"time"			
        pp =html.select("#ContentRegion .content p")                          #要换pp
        file=docx.Document()    #生成word 文档
        file.add_paragraph(ti)   #文档加入标题
        if pp:
            for ii in pp:
                file.add_paragraph(ii.text.strip())
            try:
                aurl= "E:/down/zjh/"+tim+ti+".docx" #	
                file.save(aurl.replace("\n","").replace("\t","")) #保存文
            except:
                file.save("E:/down/zjh/"+ti[:5]+".docx") #保存文
        print(index,"个详细页，等1秒")


# alist = []
with open("zjh_url.txt") as ff:
    fl= [ii.replace("\n","") for ii in ff.readlines()]
alist1= set(fl)          #url 集合	
alist = list(alist1)[460:]         #url 集合	

# for i in range(1,61):
    # geturl(bb,alist)
    # try:
        # bb.find_element_by_link_text('下一页').click()
        # time.sleep(2)
    # except:
        # print("下一页出错")


# print("写入url,")	    
# url_to(set(alist))
print(len(alist),"  个url完成，开始写入word文档")
tlist_to(alist,zlbb=bb)










